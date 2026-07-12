"""Gera documentação B2B para parceiros em formato .docx."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "API-Parceiros-B2B.docx"


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="CodeBlock")
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    if "CodeBlock" not in [style.name for style in doc.styles]:
        code_style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading("Plano Ideal", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Documentação da API B2B para Parceiros").bold = True

    version = doc.add_paragraph("Versão 1.0.0 · Julho/2026")
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    doc.add_heading("1. Visão Geral", level=1)
    doc.add_paragraph(
        "A Plano Ideal oferece uma API REST (API as a Service) para que parceiros "
        "integradores (Allvo, Record, Inova MG, entre outros) consultem viabilidade "
        "de fibra por CEP e realizem análise de crédito de forma programática, sem "
        "necessidade de acesso ao painel interno."
    )
    doc.add_paragraph(
        "Esta documentação descreve autenticação, escopos, limites, endpoints, "
        "fluxos de integração e códigos de erro. A especificação interativa "
        "(Swagger/OpenAPI) está disponível online."
    )

    doc.add_heading("1.1 URLs de referência", level=2)
    add_table(
        doc,
        ["Recurso", "URL"],
        [
            ["API Produção", "https://plano-ideal-api-production.up.railway.app"],
            [
                "Swagger UI (documentação interativa)",
                "https://plano-ideal-api-production.up.railway.app/api/docs",
            ],
            [
                "OpenAPI JSON",
                "https://plano-ideal-api-production.up.railway.app/api/docs/openapi.json",
            ],
            [
                "Painel interno (gestão de parceiros)",
                "https://plano-ideal-production.up.railway.app/interno/painel",
            ],
        ],
    )

    doc.add_heading("2. Onboarding do Parceiro", level=1)
    doc.add_paragraph(
        "O fluxo de habilitação é gerenciado pela equipe Plano Ideal (perfil administrador):"
    )
    doc.add_paragraph(
        'Cadastro do parceiro na aba "API / Parceiros" do painel interno.',
        style="List Bullet",
    )
    doc.add_paragraph(
        "Geração de uma API Key com nome descritivo (ex.: Produção, Homologação).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Seleção dos escopos necessários para a integração.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Entrega da chave ao parceiro em visualização única (write-only) — "
        "não é possível recuperá-la depois.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Compartilhamento do link do Swagger para homologação e testes.",
        style="List Bullet",
    )

    important = doc.add_paragraph()
    important.add_run("Importante:").bold = True
    doc.add_paragraph(
        "A chave segue o formato pk_live_<segredo> e deve ser armazenada com segurança.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Chaves revogadas deixam de funcionar imediatamente.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Cada consulta de crédito fica isolada pela API Key — parceiros não acessam "
        "protocolos de terceiros.",
        style="List Bullet",
    )

    doc.add_heading("3. Autenticação", level=1)
    doc.add_paragraph(
        "Todas as rotas externas exigem API Key válida, ativa e com o escopo adequado."
    )

    doc.add_heading("3.1 Formas aceitas", level=2)
    doc.add_paragraph("Header X-API-Key: pk_live_...", style="List Bullet")
    doc.add_paragraph(
        "Header Authorization: Bearer pk_live_...", style="List Bullet"
    )

    doc.add_heading("3.2 Exemplo (cURL)", level=2)
    add_code(
        doc,
        'curl -H "X-API-Key: pk_live_SEU_SEGREDO" \\\n'
        '  "https://plano-ideal-api-production.up.railway.app/api/v1/external/coverage/30130010"',
    )

    doc.add_heading("4. Escopos (Permissões)", level=1)
    add_table(
        doc,
        ["Escopo técnico", "Nome amigável", "Descrição"],
        [
            ["coverage", "Viabilidade (CEP)", "Consulta de cobertura de fibra por CEP"],
            [
                "credit",
                "Crédito (CPF/CNPJ)",
                "Consulta assíncrona de análise de crédito PAP",
            ],
        ],
    )
    doc.add_paragraph(
        "Requisições sem o escopo necessário retornam HTTP 403 com code: SCOPE_FORBIDDEN."
    )

    doc.add_heading("5. Limites de Requisição (Rate Limit)", level=1)
    add_table(
        doc,
        ["Recurso", "Limite padrão", "Janela"],
        [
            ["Viabilidade (coverage)", "60 requisições", "por minuto / API Key"],
            ["Crédito (credit)", "10 requisições", "por minuto / API Key"],
        ],
    )
    doc.add_paragraph("Ao exceder o limite, a API retorna HTTP 429 com:")
    doc.add_paragraph("code: RATE_LIMITED", style="List Bullet")
    doc.add_paragraph("retryAfterSec: segundos até nova tentativa", style="List Bullet")
    doc.add_paragraph("Header Retry-After (quando aplicável)", style="List Bullet")

    doc.add_heading("6. Endpoints", level=1)

    doc.add_heading("6.1 GET /api/v1/external/coverage/{cep}", level=2)
    doc.add_paragraph(
        "Consulta viabilidade de fibra por CEP. Retorna DTO sanitizado "
        "(sem dados internos como row_data ou source_file)."
    )
    scope = doc.add_paragraph()
    scope.add_run("Escopo necessário: coverage").bold = True
    doc.add_paragraph("Parâmetros:")
    doc.add_paragraph(
        "cep (path): 8 dígitos, com ou sem hífen (ex.: 30130-010 ou 30130010)",
        style="List Bullet",
    )
    doc.add_paragraph("Resposta 200 (exemplo):")
    add_code(
        doc,
        '{\n'
        '  "cep": "30130010",\n'
        '  "hasCoverage": true,\n'
        '  "operators": [\n'
        '    {\n'
        '      "name": "Nio",\n'
        '      "mode": "facades",\n'
        '      "items": ["7", "15", "39 +2"]\n'
        "    },\n"
        '    {\n'
        '      "name": "Vero",\n'
        '      "mode": "streets",\n'
        '      "items": ["RUA EXEMPLO - CENTRO"]\n'
        "    }\n"
        "  ]\n"
        "}",
    )

    doc.add_heading("6.2 POST /api/v1/external/credit/consult", level=2)
    doc.add_paragraph(
        "Enfileira consulta de crédito assíncrona. Retorna HTTP 202 Accepted com protocolo (id)."
    )
    scope = doc.add_paragraph()
    scope.add_run("Escopo necessário: credit").bold = True
    doc.add_paragraph("Body (JSON):")
    add_code(
        doc,
        '{\n  "document": "12345678901",\n  "cpfRepresentative": null\n}',
    )
    doc.add_paragraph("Resposta 202 (exemplo):")
    add_code(
        doc,
        '{\n'
        '  "consultation": {\n'
        '    "id": 105,\n'
        '    "status": "queued",\n'
        '    "createdAt": "2026-07-11T22:00:00.000Z"\n'
        "  }\n"
        "}",
    )

    doc.add_heading("6.3 GET /api/v1/external/credit/consultations/{id}", level=2)
    doc.add_paragraph(
        "Polling do status da consulta até conclusão (isTerminal: true)."
    )
    doc.add_paragraph(
        "Estados canônicos: queued → processing → success | failed",
        style="List Bullet",
    )
    doc.add_paragraph("Resposta 200 — concluída (negada):")
    add_code(
        doc,
        '{\n'
        '  "consultation": {\n'
        '    "id": 105,\n'
        '    "documentMasked": "123.456.789-01",\n'
        '    "status": "success",\n'
        '    "isTerminal": true,\n'
        '    "approved": false,\n'
        '    "resultDetail": "Crédito negado",\n'
        '    "finishedAt": "2026-07-11T22:01:00.000Z",\n'
        '    "durationSeconds": 55\n'
        "  }\n"
        "}",
    )

    doc.add_heading("7. Fluxo Recomendado de Integração", level=1)
    doc.add_heading("7.1 Viabilidade (síncrono)", level=2)
    doc.add_paragraph(
        "1. Cliente informa CEP no canal do parceiro", style="List Bullet"
    )
    doc.add_paragraph("2. Parceiro chama GET /coverage/{cep}", style="List Bullet")
    doc.add_paragraph(
        "3. Exibe resultado com base em hasCoverage e operators", style="List Bullet"
    )

    doc.add_heading("7.2 Crédito (assíncrono com polling)", level=2)
    doc.add_paragraph(
        "1. Parceiro envia POST /credit/consult com documento", style="List Bullet"
    )
    doc.add_paragraph("2. Armazena o id retornado (protocolo)", style="List Bullet")
    doc.add_paragraph(
        "3. Executa polling em GET /credit/consultations/{id} a cada 3–5 segundos",
        style="List Bullet",
    )
    doc.add_paragraph("4. Interrompe quando isTerminal === true", style="List Bullet")
    doc.add_paragraph(
        "5. Interpreta approved e resultDetail para decisão comercial",
        style="List Bullet",
    )

    doc.add_heading("8. Códigos de Erro", level=1)
    add_table(
        doc,
        ["HTTP", "code", "Descrição"],
        [
            ["400", "INVALID_CEP", "CEP inválido (esperado 8 dígitos)"],
            ["400", "INVALID_DOCUMENT", "CPF/CNPJ inválido"],
            ["401", "INVALID_API_KEY", "API Key ausente ou inválida"],
            ["403", "SCOPE_FORBIDDEN", "API Key sem permissão para o recurso"],
            ["404", "NOT_FOUND", "Consulta de crédito não encontrada"],
            ["429", "RATE_LIMITED", "Limite de requisições excedido"],
            ["503", "SERVICE_UNAVAILABLE", "Serviço de crédito indisponível"],
            ["500", "—", "Erro interno do servidor"],
        ],
    )

    doc.add_heading("9. Suporte e Homologação", level=1)
    doc.add_paragraph(
        "Para homologação, utilize o Swagger UI com a API Key fornecida "
        '(botão "Authorize").'
    )
    doc.add_paragraph(
        "Em caso de dúvidas técnicas ou solicitação de nova chave, contate a equipe "
        "Plano Ideal responsável pelo cadastro do parceiro."
    )

    footer = doc.add_paragraph(
        "© 2026 Plano Ideal · Documento gerado a partir da especificação OpenAPI v1.0.0"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUTPUT)
    print(f"Gerado: {OUTPUT}")
    print(f"Tamanho: {OUTPUT.stat().st_size // 1024} KB")


if __name__ == "__main__":
    main()
